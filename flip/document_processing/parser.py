"""File format parsers."""

from pathlib import Path
from typing import Optional
import PyPDF2
import pdfplumber
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
import json
import csv

from flip.core.exceptions import DocumentProcessingError


class DocumentParser:
    """Parse various document formats into text."""
    
    def parse_file(self, file_path: str) -> str:
        """
        Parse a file and extract text content.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted text content
            
        Raises:
            DocumentProcessingError: If parsing fails
        """
        path = Path(file_path)
        extension = path.suffix.lower()
        
        parsers = {
            '.txt': self._parse_text,
            '.md': self._parse_text,
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.doc': self._parse_docx,
            '.html': self._parse_html,
            '.htm': self._parse_html,
            '.json': self._parse_json,
            '.csv': self._parse_csv,
        }
        
        # Code files - treat as text
        code_extensions = {'.py', '.js', '.java', '.cpp', '.c', '.h', 
                          '.ts', '.tsx', '.jsx', '.go', '.rs', '.rb'}
        if extension in code_extensions:
            return self._parse_code(file_path, extension)
        
        parser = parsers.get(extension)
        if not parser:
            raise DocumentProcessingError(f"No parser available for {extension}")
        
        try:
            return parser(file_path)
        except Exception as e:
            raise DocumentProcessingError(f"Failed to parse {file_path}: {str(e)}")
    
    def _parse_text(self, file_path: str) -> str:
        """Parse plain text or markdown file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    def _parse_pdf(self, file_path: str) -> str:
        """Parse PDF file."""
        text_parts = []
        
        try:
            # Try pdfplumber first (better for tables)
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
        except:
            # Fallback to PyPDF2
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
        
        return '\n\n'.join(text_parts)
    
    def _parse_docx(self, file_path: str) -> str:
        """Parse DOCX file."""
        doc = DocxDocument(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        return '\n\n'.join(text_parts)
    
    def _parse_html(self, file_path: str) -> str:
        """Parse HTML file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            soup = BeautifulSoup(f.read(), 'lxml')
        
        # Remove script and style elements
        for script in soup(['script', 'style']):
            script.decompose()
        
        # Get text
        text = soup.get_text()
        
        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        
        return text
    
    def _parse_json(self, file_path: str) -> str:
        """Parse JSON file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Convert to formatted string
        return json.dumps(data, indent=2)
    
    def _parse_csv(self, file_path: str) -> str:
        """Parse CSV file."""
        text_parts = []
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            reader = csv.reader(f)
            for row in reader:
                text_parts.append(' | '.join(row))
        
        return '\n'.join(text_parts)
    
    def _parse_code(self, file_path: str, extension: str) -> str:
        """Parse code file, preserving structure."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        # Add language identifier for context
        lang_map = {
            '.py': 'Python',
            '.js': 'JavaScript',
            '.ts': 'TypeScript',
            '.java': 'Java',
            '.cpp': 'C++',
            '.c': 'C',
            '.go': 'Go',
            '.rs': 'Rust',
            '.rb': 'Ruby',
        }
        
        language = lang_map.get(extension, 'Code')
        return f"[{language} Code]\n\n{content}"
