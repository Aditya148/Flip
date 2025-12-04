"""DOCX image extraction using python-docx."""

from typing import List
from pathlib import Path
import io

try:
    from docx import Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

from flip.document.image_extractor import (
    BaseImageExtractor,
    ExtractedImage,
    ImageMetadata
)


class DOCXImageExtractor(BaseImageExtractor):
    """Extract images from DOCX files using python-docx."""
    
    def __init__(self, max_image_size: int = 1024, min_image_size: int = 100, extract_context: bool = True):
        """
        Initialize DOCX image extractor.
        
        Args:
            max_image_size: Maximum dimension for extracted images
            min_image_size: Minimum dimension to consider
            extract_context: Whether to extract surrounding text as context
        """
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX image extraction. "
                "Install with: pip install python-docx"
            )
        
        super().__init__(max_image_size, min_image_size)
        self.extract_context = extract_context
    
    def extract_images(self, file_path: str) -> List[ExtractedImage]:
        """
        Extract images from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            List of extracted images with metadata
        """
        path = Path(file_path)
        
        if path.suffix.lower() not in ['.docx', '.doc']:
            return []
        
        extracted_images = []
        
        try:
            doc = Document(file_path)
            
            # Get all relationships that are images
            image_index = 0
            
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        # Get image data
                        image_data = rel.target_part.blob
                        
                        # Convert to PIL Image
                        image = Image.open(io.BytesIO(image_data))
                        
                        # Convert to RGB if needed
                        if image.mode not in ('RGB', 'L'):
                            image = image.convert('RGB')
                        
                        width, height = image.size
                        
                        # Filter small images
                        if not self._should_include_image(width, height):
                            continue
                        
                        # Resize if needed
                        image = self._resize_if_needed(image)
                        
                        # Extract context
                        context = None
                        caption = None
                        if self.extract_context:
                            context, caption = self._extract_context(doc, rel.rId)
                        
                        # Determine format
                        image_format = rel.target_ref.split('.')[-1].upper()
                        
                        metadata = ImageMetadata(
                            source_file=str(path),
                            page_number=None,  # DOCX doesn't have page numbers in the same way
                            image_index=image_index,
                            width=image.size[0],
                            height=image.size[1],
                            format=image_format,
                            caption=caption,
                            alt_text=None,
                            position=None,
                            context=context
                        )
                        
                        extracted_images.append(
                            ExtractedImage(image=image, metadata=metadata)
                        )
                        
                        image_index += 1
                        
                    except Exception as e:
                        print(f"Error extracting image {image_index}: {e}")
                        continue
            
        except Exception as e:
            print(f"Error processing DOCX {file_path}: {e}")
        
        return extracted_images
    
    def _extract_context(self, doc, image_rid: str, context_chars: int = 500):
        """
        Extract text context around the image.
        
        Args:
            doc: Document object
            image_rid: Relationship ID of the image
            context_chars: Number of characters to extract
            
        Returns:
            Tuple of (context text, caption)
        """
        try:
            context_paragraphs = []
            caption = None
            
            # Iterate through paragraphs to find the one with the image
            for para in doc.paragraphs:
                # Check if paragraph contains the image
                if self._paragraph_contains_image(para, image_rid):
                    # Get text from this paragraph (might be caption)
                    para_text = para.text.strip()
                    if para_text:
                        caption = para_text
                    
                    # Get surrounding paragraphs for context
                    para_index = doc.paragraphs.index(para)
                    
                    # Get previous paragraph
                    if para_index > 0:
                        prev_text = doc.paragraphs[para_index - 1].text.strip()
                        if prev_text:
                            context_paragraphs.append(prev_text)
                    
                    # Get next paragraph
                    if para_index < len(doc.paragraphs) - 1:
                        next_text = doc.paragraphs[para_index + 1].text.strip()
                        if next_text:
                            context_paragraphs.append(next_text)
                    
                    break
            
            context = " ".join(context_paragraphs)
            if len(context) > context_chars:
                context = context[:context_chars] + "..."
            
            return context if context else None, caption
            
        except Exception:
            return None, None
    
    def _paragraph_contains_image(self, paragraph, image_rid: str) -> bool:
        """Check if paragraph contains the image with given relationship ID."""
        try:
            for run in paragraph.runs:
                if hasattr(run, '_element'):
                    # Check for drawing elements
                    for drawing in run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                        # Check if this drawing references our image
                        blips = drawing.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                        for blip in blips:
                            embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                            if embed_id == image_rid:
                                return True
        except Exception:
            pass
        
        return False
    
    @classmethod
    def is_supported(cls, file_path: str) -> bool:
        """Check if file is a DOCX."""
        return Path(file_path).suffix.lower() in ['.docx', '.doc']
